from docx import Document

def save_txt_docx(txt_path, docx_path, orig_text, en_text, language="unknown"):
    # Save TXT file
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(f"Language detected: {language}\n\n")
        f.write("Original transcription:\n")
        f.write(orig_text + "\n")
        
        # Only add English translation section if it exists (for Hindi audio)
        if en_text and en_text.strip():
            f.write("\nEnglish translation:\n")
            f.write(en_text + "\n")

    # Save DOCX file
    doc = Document()
    doc.add_heading("Voice Note", level=1)
    doc.add_paragraph(f"Language detected: {language}")
    doc.add_heading("Original transcription", level=2)
    doc.add_paragraph(orig_text)
    
    # Only add English translation section if it exists (for Hindi audio)
    if en_text and en_text.strip():
        doc.add_heading("English translation", level=2)
        doc.add_paragraph(en_text)
    
    doc.save(docx_path)